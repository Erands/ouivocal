from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from faster_whisper import WhisperModel
from deep_translator import GoogleTranslator
from docx import Document
import edge_tts
import asyncio
import os, uuid

app = Flask(__name__)
CORS(app)

# ✅ SAFE MODEL INIT
model = None

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route("/")
def home():
    return "Lingora API Running 🚀"

# =========================
# TRANSLATE
# =========================
def do_translate(text, direction):
    try:
        if direction == "fr-en":
            return GoogleTranslator(source='fr', target='en').translate(text)
        else:
            return GoogleTranslator(source='en', target='fr').translate(text)
    except:
        return text


# =========================
# 🔊 EDGE TTS
# =========================
async def generate_voice(text, lang, voiceType, output_path):

    if lang == "en":
        voice = "en-US-GuyNeural" if voiceType == "male" else "en-US-JennyNeural"
    else:
        voice = "fr-FR-HenriNeural" if voiceType == "male" else "fr-FR-DeniseNeural"

    communicate = edge_tts.Communicate(text, voice)
    await communicate.save(output_path)


def create_voice(text, lang, voiceType, output_path):
    asyncio.run(generate_voice(text, lang, voiceType, output_path))


# =========================
# 🎤 AUDIO
# =========================
@app.route("/translate", methods=["POST"])
def translate_audio():
    global model

    try:
        audio = request.files["audio"]
        direction = request.form.get("direction")
        voiceType = request.form.get("voiceType", "female")

        path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.webm")
        audio.save(path)

        source_lang = "fr" if direction == "fr-en" else "en"

        # ✅ SAFE LOAD (LOW MEMORY MODE)
        if model is None:
            model = WhisperModel(
                "base",
                device="cpu",
                compute_type="int8"
            )

        segments, _ = model.transcribe(path, language=source_lang)
        text = " ".join([s.text for s in segments])

        translated = do_translate(text, direction)

        filename = f"{uuid.uuid4().hex}.mp3"
        out = os.path.join(UPLOAD_FOLDER, filename)

        create_voice(
            translated,
            "en" if direction == "fr-en" else "fr",
            voiceType,
            out
        )

        return jsonify({
            "translated": translated,
            "audio": request.host_url + f"uploads/{filename}"
        })

    except Exception as e:
        return jsonify({"error": str(e)})


# =========================
# ✍️ TEXT
# =========================
@app.route("/translate-text", methods=["POST"])
def translate_text():
    try:
        data = request.get_json()
        text = data.get("text")
        direction = data.get("direction")
        voiceType = data.get("voiceType", "female")

        translated = do_translate(text, direction)

        filename = f"{uuid.uuid4().hex}.mp3"
        out = os.path.join(UPLOAD_FOLDER, filename)

        create_voice(
            translated,
            "en" if direction == "fr-en" else "fr",
            voiceType,
            out
        )

        return jsonify({
            "translated": translated,
            "audio": request.host_url + f"uploads/{filename}"
        })

    except Exception as e:
        return jsonify({"error": str(e)})


# =========================
# 📄 DOC
# =========================
@app.route("/translate-doc", methods=["POST"])
def translate_doc():
    try:
        file = request.files["file"]
        direction = request.form.get("direction")

        if not file.filename.endswith(".docx"):
            return jsonify({"error": "Use DOCX for best format"}), 400

        doc = Document(file)
        new_doc = Document()

        for para in doc.paragraphs:
            new_para = new_doc.add_paragraph()

            for run in para.runs:
                translated = do_translate(run.text, direction)

                r = new_para.add_run(translated)
                r.bold = run.bold
                r.italic = run.italic
                r.underline = run.underline

                if run.font.size:
                    r.font.size = run.font.size

        output = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.docx")
        new_doc.save(output)

        return send_file(output, as_attachment=True, download_name="translated.docx")

    except Exception as e:
        return jsonify({"error": str(e)})


# =========================
# SERVE AUDIO
# =========================
@app.route("/uploads/<filename>")
def serve_file(filename):
    path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(path):
        return "Not Found", 404
    return send_file(path)


if __name__ == "__main__":
    app.run(debug=True)
